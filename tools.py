import os
import json
import re
from datetime import datetime, timedelta
from io import BytesIO
from typing import Dict, Any, Optional

import pandas as pd
import requests

try:
    from docx import Document
except Exception:
    Document = None

SKILLS_DIR = os.path.join(os.path.dirname(__file__), "skills")

# -----------------------
# Data / tool functions
# -----------------------

def read_table(file_obj) -> pd.DataFrame:
    """Read Excel/CSV from Streamlit upload object."""
    name = getattr(file_obj, "name", "")
    if name.lower().endswith(".csv"):
        return pd.read_csv(file_obj)
    # Excel
    return pd.read_excel(file_obj)


_CANONICAL = {
    "project": ["project", "项目", "项目名称", "proj", "工程", "project_name"],
    "module": ["module", "模块", "子系统", "workstream", "领域"],
    "task": ["task", "任务", "事项", "需求", "issue", "title"],
    "owner": ["owner", "负责人", "owner_name", "assignee", "经办人"],
    "status": ["status", "状态", "进度状态", "stage"],
    "priority": ["priority", "优先级", "p", "prio"],
    "due_date": ["due_date", "截止", "截止日期", "deadline", "due", "计划完成日期"],
    "progress": ["progress", "完成度", "percent", "百分比", "进展"],
    "blocker": ["blocker", "阻塞", "障碍", "block", "问题", "风险点"],
    "risk": ["risk", "风险", "risk_level", "风险等级"],
}

_STATUS_MAP = {
    "done": ["done", "完成", "已完成", "closed", "resolved"],
    "doing": ["doing", "进行中", "开发中", "处理中", "in progress"],
    "blocked": ["blocked", "阻塞", "卡住", "blocked by", "block"],
    "todo": ["todo", "未开始", "待办", "open", "pending"],
}

_PRIORITY_ORDER = ["P0", "P1", "P2", "P3", "高", "中", "低"]


def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Map user columns to canonical schema if possible."""
    df = df.copy()
    col_map = {}
    lower_cols = {c: str(c).strip().lower() for c in df.columns}

    # Build reverse lookup
    for canonical, aliases in _CANONICAL.items():
        for c in df.columns:
            lc = lower_cols[c]
            if lc == canonical or any(lc == str(a).strip().lower() for a in aliases):
                col_map[c] = canonical
                break

    df = df.rename(columns=col_map)

    # Ensure canonical columns exist
    for c in _CANONICAL.keys():
        if c not in df.columns:
            df[c] = None

    # Normalize status
    df["status"] = df["status"].astype(str).fillna("").apply(_normalize_status)

    # Normalize due_date
    df["due_date"] = pd.to_datetime(df["due_date"], errors="coerce")

    # Normalize progress
    df["progress"] = df["progress"].apply(_normalize_progress)

    # Normalize priority
    df["priority"] = df["priority"].astype(str).fillna("").apply(_normalize_priority)

    return df


def _normalize_status(x: str) -> str:
    s = str(x).strip().lower()
    if s in ("nan", "none"):
        return "unknown"
    for k, aliases in _STATUS_MAP.items():
        if any(s == a.lower() for a in aliases):
            return k
    # Heuristics
    if "完" in s or "close" in s:
        return "done"
    if "阻" in s or "block" in s:
        return "blocked"
    if "进行" in s or "progress" in s or "doing" in s:
        return "doing"
    if "未" in s or "todo" in s or "open" in s:
        return "todo"
    return s or "unknown"


def _normalize_progress(x) -> float:
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return float("nan")
    if isinstance(x, (int, float)):
        # 0-1 or 0-100
        if 0 <= x <= 1:
            return float(x) * 100.0
        return float(x)
    s = str(x).strip()
    if not s:
        return float("nan")
    s = s.replace("%", "")
    try:
        v = float(s)
        if 0 <= v <= 1:
            v *= 100.0
        return v
    except Exception:
        return float("nan")


def _normalize_priority(x: str) -> str:
    s = str(x).strip().upper()
    if s in ("NAN", "NONE", ""):
        return "P2"
    # Convert Chinese
    if s in ("高", "HIGH"):
        return "P0"
    if s in ("中", "MEDIUM"):
        return "P1"
    if s in ("低", "LOW"):
        return "P2"
    # Keep P0-P3
    if re.match(r"^P[0-3]$", s):
        return s
    return s


def compute_weekly_kpis(df: pd.DataFrame) -> Dict[str, Any]:
    """Compute basic KPIs for weekly report. Fully local and deterministic."""
    now = datetime.now()
    df = df.copy()

    total = len(df)
    done = int((df["status"] == "done").sum())
    doing = int((df["status"] == "doing").sum())
    blocked = int((df["status"] == "blocked").sum())
    todo = int((df["status"] == "todo").sum())

    overdue_mask = df["due_date"].notna() & (df["due_date"] < now) & (df["status"] != "done")
    overdue = int(overdue_mask.sum())

    # By owner
    by_owner = (
        df.groupby("owner", dropna=False)["status"]
          .value_counts()
          .unstack(fill_value=0)
          .reset_index()
          .fillna({"owner": "（未填写）"})
          .to_dict(orient="records")
    )

    # Top risks: blocked or overdue or explicit risk high
    df["_risk_score"] = 0
    df.loc[df["status"] == "blocked", "_risk_score"] += 3
    df.loc[overdue_mask, "_risk_score"] += 2
    df.loc[df["priority"].astype(str).str.upper().isin(["P0", "P1"]), "_risk_score"] += 1
    df.loc[df["risk"].astype(str).str.contains("高|high|P0", case=False, na=False), "_risk_score"] += 2

    top_risks = (
        df.sort_values(["_risk_score"], ascending=False)
          .head(5)[["project", "module", "task", "owner", "status", "priority", "due_date", "blocker", "risk", "_risk_score"]]
    )
    # Serialize datetime
    top_risks_records = []
    for _, r in top_risks.iterrows():
        d = r.to_dict()
        if pd.notna(d.get("due_date")):
            d["due_date"] = pd.to_datetime(d["due_date"]).strftime("%Y-%m-%d")
        top_risks_records.append(d)

    # Project summary
    by_project = (
        df.groupby("project", dropna=False)["status"]
          .value_counts()
          .unstack(fill_value=0)
          .reset_index()
          .fillna({"project": "（未填写）"})
          .to_dict(orient="records")
    )

    kpis = {
        "generated_at": now.strftime("%Y-%m-%d %H:%M"),
        "total_tasks": total,
        "done": done,
        "doing": doing,
        "blocked": blocked,
        "todo": todo,
        "overdue": overdue,
        "by_owner": by_owner,
        "by_project": by_project,
        "top_risks": top_risks_records,
    }
    return kpis


def dataframe_to_markdown_table(df: pd.DataFrame, max_rows: int = 25) -> str:
    """Compact table to feed into LLM (avoid huge context)."""
    cols = ["project", "module", "task", "owner", "status", "priority", "due_date", "progress", "blocker", "risk"]
    dff = df[cols].copy().head(max_rows)
    if "due_date" in dff.columns:
        dff["due_date"] = dff["due_date"].dt.strftime("%Y-%m-%d")
    return dff.to_markdown(index=False)


# -----------------------
# Skill pack
# -----------------------

def load_skill_pack(name: str) -> Dict[str, str]:
    """Load skill pack files from ./skills/<name>"""
    base = os.path.join(SKILLS_DIR, name)
    if not os.path.isdir(base):
        raise FileNotFoundError(f"Skill pack not found: {base}")
    def _read(fn: str) -> str:
        path = os.path.join(base, fn)
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    return {
        "name": name,
        "skill_md": _read("skill.md"),
        "template_md": _read("report_template.md"),
        "rubric": _read("rubric.yaml"),
    }


# -----------------------
# Ollama client
# -----------------------

def call_ollama_chat(host: str, model: str, system: str, user: str, temperature: float = 0.3) -> str:
    """Call local Ollama chat API. Requires Ollama running at host (default 127.0.0.1:11434)."""
    url = host.rstrip("/") + "/api/chat"
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "options": {"temperature": float(temperature)},
        "stream": False,
    }
    r = requests.post(url, json=payload, timeout=120)
    r.raise_for_status()
    data = r.json()
    # Ollama returns {"message":{"role":"assistant","content":"..."}, ...}
    return data.get("message", {}).get("content", "").strip()


def build_report_prompt(user_request: str, kpis: Dict[str, Any], table_md: str, skill_pack: Optional[Dict[str, str]]) -> Dict[str, str]:
    """Compose system/user prompts. If skill_pack is provided, enforce SOP + template output."""
    kpis_json = json.dumps(kpis, ensure_ascii=False, indent=2)

    if skill_pack is None:
        system = (
            "你是一个项目管理助理。请根据用户需求和提供的数据，输出清晰的项目周报（Markdown）。"
            "优先基于事实数据，不要编造。若数据缺失请标注“待补充”。"
        )
        user = f"""用户需求：
{user_request}

数据（汇总KPIs，JSON）：
{kpis_json}

数据（任务表前若干行，Markdown 表格）：
{table_md}

请输出一份项目周报（Markdown）。"""
        return {"system": system, "user": user}

    # Skill-enabled prompt
    system = f"""你是企业内部的“周报生成技能（Skill）”执行器。

[Skill 指令]
{skill_pack["skill_md"]}

[输出模板]
{skill_pack["template_md"]}

[质检 Rubric]
{skill_pack["rubric"]}

注意：
- 严格按模板输出，保留所有模板标题。
- 必须引用给定数据；不能编造项目、里程碑或指标。
- 对缺失字段用“待补充”标注；不要瞎猜。
- 输出必须是 Markdown（不加代码块围栏）。
"""
    user = f"""用户需求：
{user_request}

输入数据（汇总KPIs，JSON）：
{kpis_json}

输入数据（任务表前若干行，Markdown 表格）：
{table_md}

现在请按模板生成周报。"""
    return {"system": system, "user": user}


# -----------------------
# DOCX export (simple)
# -----------------------

def make_docx_from_markdown_text(md: str) -> bytes:
    """Very simple DOCX: store markdown as plain text. Good enough for a demo."""
    if Document is None:
        # fallback: return as text bytes (not a real docx)
        return md.encode("utf-8")

    doc = Document()
    for line in md.splitlines():
        line = line.rstrip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
